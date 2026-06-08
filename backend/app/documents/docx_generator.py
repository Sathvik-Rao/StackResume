"""DOCX resume + cover-letter generators using python-docx.

Mirrors the visual hierarchy of the PDF templates (classic_ats / modern_clean /
executive_dark) so the same resume looks consistent across formats.

Design parity with `pdf/generator.py`:
- Name large, centred, template-coloured
- Role line under it
- Contact bits joined with bullet separators, links inline
- Section headings UPPERCASE with a hairline underline rule
- Bullets with hanging indent
- Two-column heading rows for job title/dates and company/location
"""
from __future__ import annotations

import io
import os
import uuid
import zipfile
from datetime import date
from typing import Literal

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TemplateName = Literal[
    "classic_ats", "modern_clean", "executive_dark", "dark_theme", "latex_serif"
]
FontSize = Literal["small", "normal", "large"]


PALETTES_HEX = {
    "classic_ats": {
        "name": "000000", "title": "222222", "header": "000000",
        "rule": "000000", "body": "1a1a1a", "muted": "555555", "accent": "000000",
    },
    "modern_clean": {
        "name": "0b1437", "title": "3b4dc8", "header": "0b1437",
        "rule": "3b4dc8", "body": "1f2540", "muted": "5a6079", "accent": "3b4dc8",
    },
    "executive_dark": {
        "name": "0f172a", "title": "475569", "header": "0f172a",
        "rule": "94a3b8", "body": "1e293b", "muted": "64748b", "accent": "334155",
    },
    "dark_theme": {
        "bg": "0D0F1A",
        "name": "e8eaf2", "title": "a4a8ff", "header": "e8eaf2",
        "rule": "7c83ff", "body": "d0d4e8", "muted": "9098b3", "accent": "a4a8ff",
    },
    # latex_serif is a PDF-only serif layout; for DOCX we fall back to a clean
    # all-black light palette so the same selection still exports sensibly.
    "latex_serif": {
        "name": "000000", "title": "1a1a1a", "header": "000000",
        "rule": "000000", "body": "111111", "muted": "333333", "accent": "000000",
    },
}

FONT_SIZES = {"small": 9.5, "normal": 10.5, "large": 11.5}

# Matches `_ITEM_GAP` in pdf/generator/_story.py — the vertical breathing room
# between two consecutive items inside a section (two jobs, two pubs, …).
# Keep these in sync so the same resume looks the same across PDF/DOCX/ODT.
_ITEM_GAP_PT = 4


def _safe(v) -> str:
    if v is None:
        return ""
    if isinstance(v, str):
        return v.strip()
    if isinstance(v, (int, float)):
        return str(v)
    if isinstance(v, list):
        # Defensive: undo upstream `list(<string>)` that turned a sentence into
        # ["A","c","t","i","v","i","t","i","e","s",...].
        if len(v) >= 4 and all(isinstance(x, str) for x in v):
            singles = sum(1 for s in v if len(s) <= 1)
            if singles / len(v) >= 0.8:
                joined = "".join(v).strip()
                if joined:
                    return joined
        return ", ".join(_safe(x) for x in v if _safe(x))
    if isinstance(v, dict):
        parts = [
            v.get("action") or "",
            v.get("result") or "",
            v.get("metric") or "",
        ]
        joined = " ".join(p.strip() for p in parts if p.strip())
        if joined:
            return joined
        return _safe(v.get("name") or v.get("title") or v.get("value") or v.get("text") or v.get("description") or "")
    return str(v)


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str)


def _normalize_url(s: str) -> str:
    s = (s or "").strip()
    if not s:
        return ""
    if s.startswith("http://") or s.startswith("https://") or s.startswith("mailto:"):
        return s
    if "@" in s and "://" not in s:
        return f"mailto:{s}"
    return f"https://{s}"


def _set_run_font(run, size: float, color_hex: str, bold: bool = False, italic: bool = False):
    run.font.name = "Helvetica"
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color_hex)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), "Helvetica")


def _add_hyperlink(paragraph, url: str, text: str, color_hex: str, size: float):
    """Insert a true clickable hyperlink into a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Helvetica")
    rFonts.set(qn("w:hAnsi"), "Helvetica")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_paragraph_spacing(p, before: float = 0, after: float = 0, line: float | None = None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def _strip_table_borders(table):
    """Remove every border (table-level and cell-level) so the table renders
    as an invisible layout grid  -  matching the PDF, which uses borderless
    flowable tables.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else None
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Table-level borders → nil
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    # Replace any existing tblBorders
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    # Cell-level borders → nil (some Word templates draw cell-level borders
    # even when table-level borders are nil)
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(existing)
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)


def _bottom_border(paragraph, color_hex: str, size: int = 6):
    """Add a hairline rule under the paragraph (used for section dividers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _section(doc: Document, title: str, palette: dict, base: float):
    p = doc.add_paragraph()
    # `after=4` mirrors the 4pt gap the PDF generator leaves under its
    # HRFlowable rule (and the ODT section's `space_after=4`), so the gap
    # between a section heading and its first item is identical in all three
    # formats.
    _set_paragraph_spacing(p, before=10, after=4, line=1.15)
    run = p.add_run(title.upper())
    _set_run_font(run, base + 0.5, palette["header"], bold=True)
    _bottom_border(p, palette["rule"])


def _item_gap(doc: Document):
    """Append an invisible, low-height paragraph that creates the same vertical
    breathing room PDF gets from ``Spacer(1, _ITEM_GAP)`` — used between two
    consecutive items inside any list-style section."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run("")
    # Tiny font so the empty paragraph contributes only ~_ITEM_GAP_PT pt of
    # line height; matches the PDF spacer's 4pt visual gap.
    _set_run_font(r, _ITEM_GAP_PT, "FFFFFF")


def _add_contact_line(doc: Document, contact_bits: list[tuple[str, str | None]], palette: dict, base: float):
    if not contact_bits:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, after=2, line=1.4)
    # Slightly smaller than the PDF (base - 1.5)  -  Word/LibreOffice render
    # system Helvetica ~5% wider than ReportLab's embedded Helvetica metrics,
    # so the same string overflows at the same nominal size. Compensate here.
    csize = base - 2.0
    for i, (text, url) in enumerate(contact_bits):
        if i > 0:
            sep = p.add_run("  ·  ")
            _set_run_font(sep, csize, "9aa0b4")
        if url:
            _add_hyperlink(p, url, text, palette["accent"], csize)
        else:
            r = p.add_run(text)
            _set_run_font(r, csize, palette["muted"])


def _two_col_row(doc: Document, left_text: str, right_text: str,
                 left_style: dict, right_style: dict):
    """Render a left/right row using a borderless 2-column table."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # 65% / 35% split mirrors the PDF's `_two_col_head` layout.
    table.columns[0].width = Cm(11.5)
    table.columns[1].width = Cm(6.0)
    cells = table.rows[0].cells
    cells[0].width = Cm(11.5)
    cells[1].width = Cm(6.0)

    for cell in cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), "0")
            m.set(qn("w:type"), "dxa")
            tc_mar.append(m)
        tc_pr.append(tc_mar)

    lp = cells[0].paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(lp, after=0, line=1.2)
    lr = lp.add_run(left_text)
    _set_run_font(lr, left_style["size"], left_style["color"], bold=left_style.get("bold", False))

    rp = cells[1].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(rp, after=0, line=1.3)
    rr = rp.add_run(right_text)
    _set_run_font(rr, right_style["size"], right_style["color"])

    _strip_table_borders(table)


def _two_col_row_custom(doc: Document, fill_left, fill_right, after_pt: float = 0):
    """Borderless 2-col row where each cell is populated by a callback.

    Use this when the left side needs more than a single styled run
    (e.g. a hyperlink + trailing muted text). Each callback receives the
    paragraph it should append runs to.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(11.5)
    table.columns[1].width = Cm(6.0)
    cells = table.rows[0].cells
    cells[0].width = Cm(11.5)
    cells[1].width = Cm(6.0)

    for cell in cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), "0")
            m.set(qn("w:type"), "dxa")
            tc_mar.append(m)
        tc_pr.append(tc_mar)

    lp = cells[0].paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(lp, after=after_pt, line=1.2)
    fill_left(lp)

    rp = cells[1].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(rp, after=after_pt, line=1.2)
    fill_right(rp)

    _strip_table_borders(table)


_DOCX_BULLET_COL_CM = 0.39  # ~11pt  -  matches PDF bullet column width


def _bullet(doc: Document, text: str, palette: dict, base: float):
    # Two-cell table: col-0=•, col-1=text.
    # Continuation lines align perfectly regardless of font substitution.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Right column fills remaining page width (page - margins - bullet col)
    page_w_cm = 21.59  # A4/Letter ≈ 21.59cm; actual margin is 1.524cm each side
    text_col_cm = page_w_cm - 2 * 1.524 - _DOCX_BULLET_COL_CM
    bullet_w = Cm(_DOCX_BULLET_COL_CM)
    text_w = Cm(text_col_cm)
    table.columns[0].width = bullet_w
    table.columns[1].width = text_w
    # python-docx sets tblGrid widths but not cell w:tcW — set both to avoid equal-split
    table.rows[0].cells[0].width = bullet_w
    table.rows[0].cells[1].width = text_w

    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), "0")
            m.set(qn("w:type"), "dxa")
            tc_mar.append(m)
        tc_pr.append(tc_mar)

    mark_p = table.rows[0].cells[0].paragraphs[0]
    _set_paragraph_spacing(mark_p, after=2, line=1.45)
    r = mark_p.add_run("•")
    _set_run_font(r, base - 0.5, palette["body"])

    text_p = table.rows[0].cells[1].paragraphs[0]
    text_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(text_p, after=2, line=1.45)
    r2 = text_p.add_run(text)
    _set_run_font(r2, base - 0.5, palette["body"])

    _strip_table_borders(table)


def _set_doc_margins(doc: Document, *, side_cm: float, vert_cm: float):
    for section in doc.sections:
        section.left_margin = Cm(side_cm)
        section.right_margin = Cm(side_cm)
        section.top_margin = Cm(vert_cm)
        section.bottom_margin = Cm(vert_cm)


# ═══════════════════════════════════════════════════════════════════════════
# latex_serif — DOCX twin of the PDF's LaTeX layout.
#
# Same Computer Modern typeface (embedded so it renders without the font being
# installed), same Font Awesome contact icons, small-caps name, Title-Case ruled
# headers, "Company - Title / dates" hfill rows, italic technologies line, and
# itemize-indented bullets.
# ═══════════════════════════════════════════════════════════════════════════

from docx.enum.text import WD_TAB_ALIGNMENT  # noqa: E402
from docx.shared import Inches  # noqa: E402

_FONTS_DIR = os.path.join(os.path.dirname(__file__), "fonts")
_CMU = "CMU Serif"
# The bundled FA faces were renamed to clean, unique families with a "Regular"
# subfamily so Word/LibreOffice resolve them unambiguously (the stock names use
# a non-standard "Solid" subfamily that resolves to a missing Regular face).
_FA_SOLID = "SR FA Solid"
_FA_BRANDS = "SR FA Brands"

# family -> [(filename, embed-kind)] for the package font embedding.
_EMBED_FONTS = [
    (_CMU, [
        ("CMUSerif-Roman.ttf", "embedRegular"),
        ("CMUSerif-Bold.ttf", "embedBold"),
        ("CMUSerif-Italic.ttf", "embedItalic"),
        ("CMUSerif-BoldItalic.ttf", "embedBoldItalic"),
    ]),
    (_FA_SOLID, [("fa-solid-900.ttf", "embedRegular")]),
    (_FA_BRANDS, [("fa-brands-400.ttf", "embedRegular")]),
]

# contact-field -> (font family, FontAwesome glyph)
_FA_ICONS = {
    "phone":    (_FA_SOLID, ""),
    "email":    (_FA_SOLID, ""),
    "location": (_FA_SOLID, ""),
    "linkedin": (_FA_BRANDS, ""),
    "github":   (_FA_BRANDS, ""),
    "web":      (_FA_SOLID, ""),
}

# Geometry mirroring the PDF (0.4in side margins on US-Letter, itemize indent).
_SERIF_MARGIN_IN = 0.4
_SERIF_CONTENT_CM = (8.5 - 2 * _SERIF_MARGIN_IN) * 2.54          # 19.558cm text width
_SERIF_BULLET_POS_CM = 14 / 72 * 2.54                            # bullet glyph @ 14pt
_SERIF_BULLET_TEXT_CM = 25 / 72 * 2.54                           # text @ 25pt


def _obfuscate(data: bytes, guid_hex: str) -> bytes:
    """OOXML font obfuscation: XOR the first 32 bytes with the reversed GUID."""
    key = bytes.fromhex(guid_hex)[::-1]
    out = bytearray(data)
    for i in range(32):
        out[i] ^= key[i % 16]
    return bytes(out)


def _embed_fonts_docx(docx_bytes: bytes) -> bytes:
    """Embed the bundled CMU Serif + Font Awesome faces into a saved .docx so the
    LaTeX look survives on machines without those fonts installed."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    zin = zipfile.ZipFile(io.BytesIO(docx_bytes))
    items: dict[str, bytes] = {n: zin.read(n) for n in zin.namelist()}

    # Obfuscate the font binaries and collect <w:font> / relationship entries.
    font_xml, rels_xml, fidx = "", "", 0
    for family, faces in _EMBED_FONTS:
        font_xml += f'<w:font w:name="{family}">'
        for fn, kind in faces:
            raw = open(os.path.join(_FONTS_DIR, fn), "rb").read()
            g = uuid.uuid4()
            fidx += 1
            items[f"word/fonts/font{fidx}.odttf"] = _obfuscate(raw, g.hex)
            rid = f"rIdFont{fidx}"
            font_xml += (f'<w:{kind} r:id="{rid}" w:fontKey="{{{str(g).upper()}}}" '
                         f'w:subsetted="0"/>')
            rels_xml += (f'<Relationship Id="{rid}" Type="{R}/font" '
                         f'Target="fonts/font{fidx}.odttf"/>')
        font_xml += "</w:font>"

    # word/fontTable.xml — merge into the existing table, or create one.
    ft = items.get("word/fontTable.xml")
    if ft and b"</w:fonts>" in ft:
        items["word/fontTable.xml"] = ft.replace(
            b"</w:fonts>", font_xml.encode() + b"</w:fonts>")
    else:
        items["word/fontTable.xml"] = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:fonts xmlns:w="{W}" xmlns:r="{R}">{font_xml}</w:fonts>'
        ).encode()

    # word/_rels/fontTable.xml.rels — merge or create.
    fr = items.get("word/_rels/fontTable.xml.rels")
    if fr and b"</Relationships>" in fr:
        items["word/_rels/fontTable.xml.rels"] = fr.replace(
            b"</Relationships>", rels_xml.encode() + b"</Relationships>")
    else:
        items["word/_rels/fontTable.xml.rels"] = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/'
            f'relationships">{rels_xml}</Relationships>'
        ).encode()

    # [Content_Types].xml — declare the obfuscated-font extension (+ fontTable).
    ct = items["[Content_Types].xml"].decode()
    if "obfuscatedFont" not in ct:
        ct = ct.replace("</Types>",
            '<Default Extension="odttf" ContentType='
            '"application/vnd.openxmlformats-officedocument.obfuscatedFont"/></Types>')
    if "/word/fontTable.xml" not in ct:
        ct = ct.replace("</Types>",
            '<Override PartName="/word/fontTable.xml" ContentType='
            '"application/vnd.openxmlformats-officedocument.wordprocessingml.'
            'fontTable+xml"/></Types>')
    items["[Content_Types].xml"] = ct.encode()

    # word/_rels/document.xml.rels — reference the fontTable part.
    dr = items.get("word/_rels/document.xml.rels", b"")
    if b"fontTable" not in dr and b"</Relationships>" in dr:
        items["word/_rels/document.xml.rels"] = dr.replace(
            b"</Relationships>",
            f'<Relationship Id="rIdFontTable" Type="{R}/fontTable" '
            f'Target="fontTable.xml"/></Relationships>'.encode())

    # word/settings.xml — turn on embedded-font usage.
    st = items.get("word/settings.xml", b"").decode()
    if st and "embedTrueTypeFonts" not in st:
        idx = st.find(">", st.find("<w:settings")) + 1
        st = (st[:idx] + "<w:embedTrueTypeFonts/><w:embedSystemFonts/>"
              '<w:saveSubsetFonts w:val="false"/>' + st[idx:])
        items["word/settings.xml"] = st.encode()

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, d in items.items():
            zout.writestr(n, d)
    return out.getvalue()


def _srun(p, text: str, size: float, color: str, *, bold=False, italic=False,
          font: str = _CMU, underline=False):
    """Add a run in an explicit font (defaults to CMU Serif)."""
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    run.bold = bold
    run.italic = italic
    if underline:
        run.font.underline = True
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), font)
    rPr.insert(0, rFonts)
    return run


def _sicon(p, kind: str, color: str, size: float):
    spec = _FA_ICONS.get(kind)
    if not spec:
        return
    face, glyph = spec
    _srun(p, glyph, size, color, font=face)
    _srun(p, " ", size, color)  # nbsp gap to the text


def _ssection(doc: Document, title: str, palette: dict, base: float):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=6, after=3, line=1.05)
    _srun(p, title, base + 2.5, palette["header"], bold=True)
    _bottom_border(p, palette["rule"], size=4)


def _shrow(doc: Document, fill_left, right_text: str, right_size: float,
           right_color: str, *, right_bold=True, right_italic=False, after=0):
    """LaTeX \\hfill row via a right-aligned tab stop: left text flows, the right
    item (date/location) hugs the right margin on the same line."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, after=after, line=1.12)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(_SERIF_CONTENT_CM), WD_TAB_ALIGNMENT.RIGHT)
    fill_left(p)
    if right_text:
        p.add_run("\t")
        _srun(p, right_text, right_size, right_color, bold=right_bold, italic=right_italic)
    return p


def _sbullet(doc: Document, palette: dict, base: float, fill_text=None, text: str | None = None):
    """Itemize-style bullet: glyph indented under the entry, hanging text."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(_SERIF_BULLET_TEXT_CM)
    pf.first_line_indent = Cm(-(_SERIF_BULLET_TEXT_CM - _SERIF_BULLET_POS_CM))
    _set_paragraph_spacing(p, after=1, line=1.18)
    pf.tab_stops.add_tab_stop(Cm(_SERIF_BULLET_TEXT_CM), WD_TAB_ALIGNMENT.LEFT)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _srun(p, "•", base - 0.5, palette["body"])
    p.add_run("\t")
    if fill_text is not None:
        fill_text(p)
    else:
        _srun(p, text or "", base - 0.5, palette["body"])


def _build_resume_doc_serif(resume: dict, base: float) -> Document:
    """LaTeX-style DOCX (Computer Modern + Font Awesome) matching the PDF."""
    palette = PALETTES_HEX["latex_serif"]
    body, muted, accent = palette["body"], palette["muted"], palette["accent"]
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(_SERIF_MARGIN_IN)
        section.right_margin = Inches(_SERIF_MARGIN_IN)
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.4)

    pi = resume.get("personal_info", {}) or {}

    # ── Header: small-caps name, optional title, icon contact line ──────────
    name = _safe(pi.get("full_name")) or "Your Name"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, after=1, line=1.0)
    nr = _srun(p, name, base + 12, palette["name"], bold=True)
    nr.font.small_caps = True

    role = _safe(pi.get("professional_title"))
    if role:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, after=2, line=1.2)
        _srun(p, role, base + 0.5, muted, italic=True)

    # Only emit a contact line when there's at least one contact field, so a
    # bare resume doesn't leave an empty centred paragraph (matches the ODT).
    if any(_safe(pi.get(k)) for k in
           ("phone", "email", "location", "linkedin", "github", "website", "portfolio")):
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(cp, after=2, line=1.3)
        csize = base - 0.5
        first = True

        def _sep():
            nonlocal first
            if not first:
                _srun(cp, "    ", csize, body)
            first = False

        phone = _safe(pi.get("phone"))
        if phone:
            _sep(); _sicon(cp, "phone", body, csize); _srun(cp, phone, csize, body)
        email = _safe(pi.get("email"))
        if email:
            _sep(); _sicon(cp, "email", body, csize)
            _add_hyperlink_serif(cp, _normalize_url(email), email, accent, csize)
        location = _safe(pi.get("location"))
        if location:
            _sep(); _sicon(cp, "location", body, csize); _srun(cp, location, csize, body)
        for key in ("linkedin", "github", "website", "portfolio"):
            v = _safe(pi.get(key))
            if v:
                display = v.replace("https://", "").replace("http://", "").rstrip("/")
                _sep(); _sicon(cp, key if key in ("linkedin", "github") else "web", body, csize)
                _add_hyperlink_serif(cp, _normalize_url(v), display, accent, csize)

    # ── Summary ─────────────────────────────────────────────────────────────
    summary = _safe(resume.get("professional_summary", ""))
    if summary:
        _ssection(doc, "Summary", palette, base)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_paragraph_spacing(p, after=1, line=1.2)
        _srun(p, summary, base, body)

    # ── Core Competencies (3-col grid) ──────────────────────────────────────
    competencies = [_safe(c) for c in (resume.get("core_competencies") or []) if _safe(c)]
    if competencies:
        _ssection(doc, "Core Competencies", palette, base)
        cols = 3
        rows = (len(competencies) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = True
        for i, item in enumerate(competencies):
            cell = table.rows[i % rows].cells[i // rows]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), "0"); m.set(qn("w:type"), "dxa"); tc_mar.append(m)
            tc_pr.append(tc_mar)
            cpp = cell.paragraphs[0]
            _set_paragraph_spacing(cpp, after=1.5, line=1.2)
            _srun(cpp, f"•  {item}", base - 0.5, body)
        _strip_table_borders(table)

    # ── Professional Experience ─────────────────────────────────────────────
    experience = resume.get("experience") or []
    if experience:
        _ssection(doc, "Professional Experience", palette, base)
        for exp in experience:
            title_str = _safe(exp.get("title"))
            company_str = _safe(exp.get("company"))
            loc_str = _safe(exp.get("location"))
            start = _safe(exp.get("start_date"))
            end = _safe(exp.get("end_date")) or ("Present" if exp.get("current") else "Present")
            emp_type = _safe(exp.get("employment_type", ""))
            head_left = " – ".join([x for x in (company_str, title_str) if x]) or title_str
            if emp_type and emp_type.lower() not in ("full-time", "fulltime", ""):
                head_left = f"{head_left} ({emp_type})"
            _shrow(doc, lambda p, t=head_left: _srun(p, t, base, body, bold=True),
                   _serif_dates(start, end), base, body, right_bold=True)

            techs = [_safe(t) for t in (exp.get("technologies") or []) if _safe(t)]
            team = _safe(exp.get("team_size"))
            if techs or team or loc_str:
                def _tech(p, _techs=techs, _team=team):
                    if _techs:
                        _srun(p, "Technologies", base - 0.5, body, bold=True, italic=True)
                        _srun(p, f": {', '.join(_techs)}", base - 0.5, body, italic=True)
                        if _team:
                            _srun(p, f"  ·  Team: {_team}", base - 0.5, body, italic=True)
                    elif _team:
                        _srun(p, f"Team: {_team}", base - 0.5, body, italic=True)
                _shrow(doc, _tech, loc_str, base - 0.5, body, right_bold=False, right_italic=True)

            bullets = list(exp.get("responsibilities") or []) + list(exp.get("achievements") or [])
            seen: set[str] = set()
            for b in bullets:
                b_str = _safe(b)
                if b_str and b_str.lower() not in seen:
                    seen.add(b_str.lower())
                    _sbullet(doc, palette, base, text=b_str)
            _serif_gap(doc)

    # ── Projects ────────────────────────────────────────────────────────────
    projects = resume.get("projects") or []
    if projects:
        _ssection(doc, "Projects", palette, base)
        for proj in projects:
            p_name = _safe(proj.get("name"))
            url_raw = _safe(proj.get("url") or proj.get("github") or "")
            url = _normalize_url(url_raw) if url_raw else None
            tech = ", ".join(_safe(t) for t in (proj.get("technologies") or []) if _safe(t))
            proj_role = _safe(proj.get("role"))
            proj_type = _safe(proj.get("type"))
            name_lc = p_name.lower()
            descriptor = next((d for d in (proj_role, proj_type) if d and d.lower() not in name_lc), "")
            dates_str = _serif_dates(_safe(proj.get("start_date")), _safe(proj.get("end_date")))

            def _proj_head(p, _n=p_name, _u=url, _d=descriptor, _t=tech):
                if _u:
                    _add_hyperlink_serif(p, _u, _n, body, base, bold=True, underline=False)
                else:
                    _srun(p, _n, base, body, bold=True)
                if _d:
                    _srun(p, f" ({_d})", base, body, bold=True)
                if _t:
                    _srun(p, " | ", base, body)
                    _srun(p, _t, base - 0.5, body, italic=True)
            _shrow(doc, _proj_head, dates_str, base, body, right_bold=True)
            desc = _safe(proj.get("description"))
            if desc:
                _sbullet(doc, palette, base, text=desc)
            for h in (proj.get("highlights") or []):
                h_str = _safe(h)
                if h_str:
                    _sbullet(doc, palette, base, text=h_str)
            _serif_gap(doc)

    # ── Education ───────────────────────────────────────────────────────────
    education = resume.get("education") or []
    if education:
        _ssection(doc, "Education", palette, base)
        for edu in education:
            deg = _safe(edu.get("degree"))
            field = _safe(edu.get("field_of_study"))
            inst = _safe(edu.get("institution"))
            loc = _safe(edu.get("location"))
            start = _safe(edu.get("start_date"))
            end = _safe(edu.get("end_date")) or _safe(edu.get("graduation_year"))
            deg_line = deg
            if field and field.lower() not in deg.lower():
                deg_line = f"{deg} in {field}" if deg else field
            _shrow(doc, lambda p, t=inst: _srun(p, t, base, body, bold=True),
                   _serif_dates(start, end), base, body, right_bold=True)
            if deg_line or loc:
                _shrow(doc, lambda p, t=deg_line: _srun(p, t, base - 0.5, body, italic=True),
                       loc, base - 0.5, body, right_bold=False, right_italic=True)
            extras = []
            if _safe(edu.get("gpa")):
                extras.append(f"GPA: {_safe(edu.get('gpa'))}")
            if _safe(edu.get("honors")):
                extras.append(_safe(edu.get("honors")))
            cw = ", ".join(_safe(c) for c in (edu.get("relevant_coursework") or []) if _safe(c))
            if cw:
                extras.append(f"Coursework: {cw}")
            if extras:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, after=0, line=1.2)
                _srun(p, "  ·  ".join(extras), base - 1.5, muted, italic=True)
            _serif_gap(doc)

    # ── Technical Skills ────────────────────────────────────────────────────
    skills = resume.get("technical_skills") or {}
    skill_cats = [
        ("Languages", skills.get("programming_languages")),
        ("Frameworks & Libraries", skills.get("frameworks_and_libraries")),
        ("Databases", skills.get("databases")),
        ("Cloud & Infrastructure", skills.get("cloud_and_infrastructure")),
        ("DevOps & Tooling", skills.get("devops_and_tools") or skills.get("tools_and_practices")),
        ("Testing", skills.get("testing")),
        ("Methodologies", skills.get("methodologies")),
        ("Soft Skills", skills.get("soft_skills")),
    ]
    skill_rows = [(k, ", ".join(_safe(x) for x in v if _safe(x))) for k, v in skill_cats if v]
    if skill_rows:
        _ssection(doc, "Technical Skills", palette, base)
        for label, vals in skill_rows:
            if not vals:
                continue
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=1.5, line=1.2)
            _srun(p, label, base - 0.5, body, bold=True)
            _srun(p, f" : {vals}", base - 0.5, body)

    _serif_tail_sections(doc, resume, palette, base)
    return doc


def _serif_dates(start: str, end: str) -> str:
    if start and end:
        return f"{start} – {end}"
    return end or start or ""


def _serif_gap(doc: Document):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=0, line=1.0)
    _srun(p, "", 3, "FFFFFF")


def _add_hyperlink_serif(paragraph, url: str, text: str, color_hex: str,
                         size: float, *, bold=False, underline=True):
    """Clickable hyperlink run in CMU Serif (optionally underlined like the PDF)."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(a), _CMU)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), color_hex); rPr.append(color)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _serif_tail_sections(doc: Document, resume: dict, palette: dict, base: float):
    """Open Source · Certifications · Publications · Patents · Awards · Volunteer
    · Languages · Interests · References — serif-styled, mirroring the PDF."""
    body, muted, accent = palette["body"], palette["muted"], palette["accent"]

    oss = resume.get("open_source_contributions") or []
    if oss:
        _ssection(doc, "Open Source", palette, base)
        for o in oss:
            proj = _safe(o.get("project")); role = _safe(o.get("role"))
            url_raw = _safe(o.get("url", "")); url = _normalize_url(url_raw) if url_raw else None
            p = doc.add_paragraph(); _set_paragraph_spacing(p, after=0, line=1.18)
            if url:
                _add_hyperlink_serif(p, url, proj, body, base, bold=True, underline=False)
            else:
                _srun(p, proj, base, body, bold=True)
            if role:
                _srun(p, f" – {role}", base, body, italic=True)
            o_desc = _safe(o.get("description", "")) or _safe(o.get("contribution", ""))
            if o_desc:
                pd = doc.add_paragraph(); _set_paragraph_spacing(pd, after=0, line=1.2)
                _srun(pd, o_desc, base - 1.5, muted, italic=True)
            for c in (o.get("contributions") or []):
                if _safe(c):
                    _sbullet(doc, palette, base, text=_safe(c))
            _serif_gap(doc)

    certs = resume.get("certifications") or []
    if certs:
        _ssection(doc, "Certifications", palette, base)
        for c in certs:
            url_raw = _safe(c.get("url", "")); url = _normalize_url(url_raw) if url_raw else None
            p = doc.add_paragraph(); _set_paragraph_spacing(p, after=1, line=1.2)
            if url:
                _add_hyperlink_serif(p, url, _safe(c.get("name")), body, base - 0.5, bold=True, underline=False)
            else:
                _srun(p, _safe(c.get("name")), base - 0.5, body, bold=True)
            if _safe(c.get("issuer")):
                _srun(p, f" – {_safe(c.get('issuer'))}", base - 0.5, body, italic=True)
            dp = []
            if _safe(c.get("date")) or _safe(c.get("year")):
                dp.append(f"issued {_safe(c.get('date')) or _safe(c.get('year'))}")
            if _safe(c.get("expiry")):
                dp.append(f"expires {_safe(c.get('expiry'))}")
            if dp:
                _srun(p, f"  ({', '.join(dp)})", base - 0.5, muted)
            _serif_gap(doc)

    pubs = resume.get("publications") or []
    if pubs:
        _ssection(doc, "Publications", palette, base)
        for pub in pubs:
            url_raw = _safe(pub.get("url", "")); url = _normalize_url(url_raw) if url_raw else None
            p = doc.add_paragraph(); _set_paragraph_spacing(p, after=1, line=1.2)
            if url:
                _add_hyperlink_serif(p, url, _safe(pub.get("title")), body, base - 0.5, bold=True, underline=False)
            else:
                _srun(p, _safe(pub.get("title")), base - 0.5, body, bold=True)
            if _safe(pub.get("venue")):
                _srun(p, f" – {_safe(pub.get('venue'))}", base - 0.5, body, italic=True)
            if _safe(pub.get("date")):
                _srun(p, f"  ({_safe(pub.get('date'))})", base - 0.5, muted)
            _serif_gap(doc)

    patents = resume.get("patents") or []
    if patents:
        _ssection(doc, "Patents", palette, base)
        for pt in patents:
            p = doc.add_paragraph(); _set_paragraph_spacing(p, after=1, line=1.2)
            _srun(p, _safe(pt.get("title")), base - 0.5, body, bold=True)
            if _safe(pt.get("patent_number")):
                _srun(p, f" – {_safe(pt.get('patent_number'))}", base - 0.5, body)
            if _safe(pt.get("date")):
                _srun(p, f"  ({_safe(pt.get('date'))})", base - 0.5, muted)
            _serif_gap(doc)

    awards = resume.get("awards_and_honors") or []
    if awards:
        _ssection(doc, "Awards & Honors", palette, base)
        for a in awards:
            if isinstance(a, dict):
                nm = _safe(a.get("name") or a.get("title"))
                issuer = _safe(a.get("issuer")); yr = _safe(a.get("year") or a.get("date"))
                v = nm + (f" – {issuer}" if issuer else "") + (f" ({yr})" if yr else "")
            else:
                v = _safe(a)
            if v:
                _sbullet(doc, palette, base, text=v)

    vols = resume.get("volunteer_experience") or []
    if vols:
        _ssection(doc, "Volunteer Experience", palette, base)
        for v in vols:
            org = _safe(v.get("organization")); role = _safe(v.get("role"))
            head_left = " – ".join([x for x in (org, role) if x])
            _shrow(doc, lambda p, t=head_left: _srun(p, t, base, body, bold=True),
                   _serif_dates(_safe(v.get("start_date")), _safe(v.get("end_date"))),
                   base, body, right_bold=True)
            if _safe(v.get("description")):
                pd = doc.add_paragraph(); _set_paragraph_spacing(pd, after=0, line=1.2)
                _srun(pd, _safe(v.get("description")), base - 1.5, muted, italic=True)
            _serif_gap(doc)

    langs = resume.get("languages") or []
    if langs:
        _ssection(doc, "Languages", palette, base)
        p = doc.add_paragraph(); _set_paragraph_spacing(p, after=1.5, line=1.2)
        for i, l in enumerate(langs):
            lang = _safe(l.get("language")); prof = _safe(l.get("proficiency"))
            if not lang:
                continue
            if i > 0:
                _srun(p, "  ·  ", base - 0.5, body)
            _srun(p, lang, base - 0.5, body, bold=True)
            if prof:
                _srun(p, f" ({prof})", base - 0.5, body)

    interests = [_safe(i) for i in (resume.get("interests") or []) if _safe(i)]
    if interests:
        _ssection(doc, "Interests", palette, base)
        p = doc.add_paragraph(); _set_paragraph_spacing(p, after=1.5, line=1.2)
        _srun(p, ", ".join(interests), base - 0.5, body)

    references = resume.get("references")
    if references:
        _ssection(doc, "References", palette, base)
        if isinstance(references, str):
            p = doc.add_paragraph(); _set_paragraph_spacing(p, after=1.5, line=1.2)
            _srun(p, _safe(references), base - 0.5, body)
        elif isinstance(references, list):
            for ref in references:
                if _safe(ref):
                    _sbullet(doc, palette, base, text=_safe(ref))


def _build_resume_doc(resume: dict, template: TemplateName, base: float) -> Document:
    palette = PALETTES_HEX.get(template, PALETTES_HEX["classic_ats"])
    doc = Document()
    # Match the PDF generator: 0.6in side, 0.55in top/bottom.
    _set_doc_margins(doc, side_cm=1.524, vert_cm=1.397)

    # ── Header ────────────────────────────────────────────────────────────
    pi = resume.get("personal_info", {}) or {}
    name = _safe(pi.get("full_name")) or "Your Name"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, after=1, line=1.05)
    r = p.add_run(name)
    _set_run_font(r, base + 11.5, palette["name"], bold=True)

    role = _safe(pi.get("professional_title"))
    if role:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, after=4, line=1.3)
        rr = p.add_run(role)
        _set_run_font(rr, base + 1, palette["title"])

    contact_bits: list[tuple[str, str | None]] = []
    for key in ("email", "phone", "location"):
        v = _safe(pi.get(key))
        if v:
            contact_bits.append((v, None))
    for key in ("linkedin", "github", "website", "portfolio"):
        v = _safe(pi.get(key))
        if v:
            display = v.replace("https://", "").replace("http://", "").rstrip("/")
            contact_bits.append((display, _normalize_url(v)))
    _add_contact_line(doc, contact_bits, palette, base)

    # ── Summary ───────────────────────────────────────────────────────────
    summary = _safe(resume.get("professional_summary", ""))
    if summary:
        _section(doc, "Summary", palette, base)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_paragraph_spacing(p, after=2, line=1.45)
        r = p.add_run(summary)
        _set_run_font(r, base - 0.5, palette["body"])

    # ── Core Competencies ────────────────────────────────────────────────
    competencies = [_safe(c) for c in (resume.get("core_competencies") or []) if _safe(c)]
    if competencies:
        _section(doc, "Core Competencies", palette, base)
        cols = 3
        rows = (len(competencies) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = True
        for i, item in enumerate(competencies):
            row_idx = i % rows
            col_idx = i // rows
            cell = table.rows[row_idx].cells[col_idx]
            # Zero cell padding so the grid hugs its text like the PDF version.
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), "0")
                m.set(qn("w:type"), "dxa")
                tc_mar.append(m)
            tc_pr.append(tc_mar)
            cp = cell.paragraphs[0]
            _set_paragraph_spacing(cp, after=1.5, line=1.4)
            r = cp.add_run(f"•  {item}")
            _set_run_font(r, base - 0.5, palette["body"])
        _strip_table_borders(table)

    # ── Experience ────────────────────────────────────────────────────────
    experience = resume.get("experience") or []
    if experience:
        _section(doc, "Professional Experience", palette, base)
        for exp in experience:
            title_str = _safe(exp.get("title"))
            company_str = _safe(exp.get("company"))
            loc_str = _safe(exp.get("location"))
            start = _safe(exp.get("start_date"))
            end = _safe(exp.get("end_date")) or ("Present" if exp.get("current") else "Present")
            emp_type = _safe(exp.get("employment_type", ""))

            _two_col_row(
                doc, title_str, f"{start}  -  {end}",
                {"size": base, "color": palette["body"], "bold": True},
                {"size": base - 1.5, "color": palette["muted"]},
            )
            company_inline = company_str
            if emp_type and emp_type.lower() not in ("full-time", "fulltime", ""):
                company_inline = f"{company_str} ({emp_type})" if company_str else emp_type
            if company_inline or loc_str:
                _two_col_row(
                    doc, company_inline, loc_str,
                    {"size": base - 0.5, "color": palette["accent"]},
                    {"size": base - 1.5, "color": palette["muted"]},
                )

            bullets = list(exp.get("responsibilities") or []) + list(exp.get("achievements") or [])
            seen: set[str] = set()
            for b in bullets:
                b_str = _safe(b)
                if not b_str or b_str.lower() in seen:
                    continue
                seen.add(b_str.lower())
                _bullet(doc, b_str, palette, base)

            team_size = _safe(exp.get("team_size"))
            techs = [_safe(t) for t in (exp.get("technologies") or []) if _safe(t)]
            meta_parts = []
            if team_size:
                meta_parts.append(f"Team: {team_size}")
            if techs:
                meta_parts.append(f"Tech: {', '.join(techs)}")
            if meta_parts:
                p = doc.add_paragraph()
                # `after=0` because `_item_gap` immediately below adds the
                # standard 4pt inter-item gap (matches PDF's Spacer behaviour).
                _set_paragraph_spacing(p, after=0, line=1.4)
                r = p.add_run("  ·  ".join(meta_parts))
                _set_run_font(r, base - 1.5, palette["muted"], italic=True)
            _item_gap(doc)

    # ── Technical Skills ──────────────────────────────────────────────────
    skills = resume.get("technical_skills") or {}
    skill_cats = [
        ("Languages", skills.get("programming_languages")),
        ("Frameworks & Libraries", skills.get("frameworks_and_libraries")),
        ("Databases", skills.get("databases")),
        ("Cloud & Infrastructure", skills.get("cloud_and_infrastructure")),
        ("DevOps & Tooling", skills.get("devops_and_tools")),
        ("Testing", skills.get("testing")),
        ("Methodologies", skills.get("methodologies")),
        ("Soft Skills", skills.get("soft_skills")),
    ]
    skill_rows = [(k, ", ".join(_safe(x) for x in v if _safe(x))) for k, v in skill_cats if v]
    if skill_rows:
        _section(doc, "Technical Skills", palette, base)
        for label, vals in skill_rows:
            if not vals:
                continue
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=2, line=1.4)
            r1 = p.add_run(f"{label}:  ")
            _set_run_font(r1, base - 0.5, palette["body"], bold=True)
            r2 = p.add_run(vals)
            _set_run_font(r2, base - 0.5, palette["body"])

    # ── Projects ──────────────────────────────────────────────────────────
    projects = resume.get("projects") or []
    if projects:
        _section(doc, "Projects", palette, base)
        for proj in projects:
            p_name = _safe(proj.get("name"))
            url_raw = _safe(proj.get("url") or proj.get("github") or "")
            url = _normalize_url(url_raw) if url_raw else None
            tech = ", ".join(_safe(t) for t in (proj.get("technologies") or []) if _safe(t))
            proj_role = _safe(proj.get("role"))
            proj_type = _safe(proj.get("type"))
            p_start = _safe(proj.get("start_date"))
            p_end = _safe(proj.get("end_date"))
            name_lc = p_name.lower()
            sub_parts = []
            if proj_role and proj_role.lower() not in name_lc:
                sub_parts.append(proj_role)
            if proj_type and proj_type.lower() not in name_lc:
                sub_parts.append(proj_type)
            if tech:
                sub_parts.append(tech)
            dates_str = ""
            if p_start or p_end:
                dates_str = f"{p_start}  -  {p_end}".strip(" -") if (p_start and p_end) else (p_start or p_end)

            def _fill_left(lp, _name=p_name, _url=url, _sub=sub_parts):
                if _url:
                    _add_hyperlink(lp, _url, _name, palette["accent"], base)
                else:
                    r = lp.add_run(_name)
                    _set_run_font(r, base, palette["body"], bold=True)
                if _sub:
                    r2 = lp.add_run(f"   -   {'  ·  '.join(_sub)}")
                    _set_run_font(r2, base - 1, palette["muted"])

            def _fill_right(rp, _dates=dates_str):
                if _dates:
                    rr = rp.add_run(_dates)
                    _set_run_font(rr, base - 1.5, palette["muted"])

            # after_pt ≈ one company-row's worth of breathing room, so the
            # gap before project bullets matches the gap below Experience
            # job titles (which have a company/location row between).
            _two_col_row_custom(doc, _fill_left, _fill_right, after_pt=6)
            desc = _safe(proj.get("description"))
            if desc:
                _bullet(doc, desc, palette, base)
            for h in (proj.get("highlights") or []):
                h_str = _safe(h)
                if h_str:
                    _bullet(doc, h_str, palette, base)
            _item_gap(doc)

    # ── Open Source ──────────────────────────────────────────────────────
    oss = resume.get("open_source_contributions") or []
    if oss:
        _section(doc, "Open Source", palette, base)
        for o in oss:
            name_o = _safe(o.get("project"))
            role = _safe(o.get("role"))
            url_raw = _safe(o.get("url", ""))
            url = _normalize_url(url_raw) if url_raw else None
            stars = _safe(o.get("stars", ""))
            o_lang = _safe(o.get("language", ""))
            o_desc = _safe(o.get("description", ""))
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=0, line=1.2)
            if url:
                _add_hyperlink(p, url, name_o, palette["accent"], base)
            else:
                r = p.add_run(name_o)
                _set_run_font(r, base, palette["body"], bold=True)
            if role:
                r2 = p.add_run(f"   -   {role}")
                _set_run_font(r2, base - 0.5, palette["body"])
            if o_lang:
                rl = p.add_run(f"  {o_lang}")
                _set_run_font(rl, base - 0.5, palette["muted"])
            if stars:
                r3 = p.add_run(f"  ★ {stars}")
                _set_run_font(r3, base - 0.5, palette["muted"])
            if o_desc:
                pd = doc.add_paragraph()
                _set_paragraph_spacing(pd, after=0, line=1.4)
                rd = pd.add_run(o_desc)
                _set_run_font(rd, base - 1.5, palette["muted"])
            for contrib in (o.get("contributions") or []):
                c_str = _safe(contrib)
                if c_str:
                    _bullet(doc, c_str, palette, base)
            _item_gap(doc)

    # ── Education ────────────────────────────────────────────────────────
    education = resume.get("education") or []
    if education:
        _section(doc, "Education", palette, base)
        for edu in education:
            deg = _safe(edu.get("degree"))
            field = _safe(edu.get("field_of_study"))
            inst = _safe(edu.get("institution"))
            loc = _safe(edu.get("location"))
            start = _safe(edu.get("start_date"))
            end = _safe(edu.get("end_date"))
            gpa = _safe(edu.get("gpa"))
            honors = _safe(edu.get("honors"))
            deg_line = deg + (f" in {field}" if field else "")
            dates = f"{start}  -  {end}" if (start and end) else (end or start or "")
            _two_col_row(
                doc, deg_line, dates,
                {"size": base, "color": palette["body"], "bold": True},
                {"size": base - 1.5, "color": palette["muted"]},
            )
            if inst or loc:
                _two_col_row(
                    doc, inst, loc,
                    {"size": base - 0.5, "color": palette["accent"]},
                    {"size": base - 1.5, "color": palette["muted"]},
                )
            extras = []
            if gpa:
                extras.append(f"GPA: {gpa}")
            if honors:
                extras.append(honors)
            cw = ", ".join(_safe(c) for c in (edu.get("relevant_coursework") or []) if _safe(c))
            if cw:
                extras.append(f"Coursework: {cw}")
            if extras:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, after=0, line=1.4)
                r = p.add_run("  ·  ".join(extras))
                _set_run_font(r, base - 1.5, palette["muted"])
            acts = ", ".join(_safe(a) for a in (edu.get("activities") or []) if _safe(a))
            if acts:
                pa = doc.add_paragraph()
                _set_paragraph_spacing(pa, after=0, line=1.4)
                ra = pa.add_run(f"Activities: {acts}")
                _set_run_font(ra, base - 1.5, palette["muted"])
            _item_gap(doc)

    # ── Certifications ───────────────────────────────────────────────────
    certs = resume.get("certifications") or []
    if certs:
        _section(doc, "Certifications", palette, base)
        for c in certs:
            c_name = _safe(c.get("name"))
            c_issuer = _safe(c.get("issuer"))
            c_date = _safe(c.get("date"))
            c_expiry = _safe(c.get("expiry"))
            c_cred_id = _safe(c.get("credential_id"))
            c_url_raw = _safe(c.get("url", ""))
            c_url = _normalize_url(c_url_raw) if c_url_raw else None
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=2, line=1.4)
            if c_url:
                _add_hyperlink(p, c_url, c_name, palette["accent"], base - 0.5)
            else:
                r = p.add_run(c_name)
                _set_run_font(r, base - 0.5, palette["body"], bold=True)
            if c_issuer:
                r2 = p.add_run(f"   -   {c_issuer}")
                _set_run_font(r2, base - 0.5, palette["body"])
            date_parts = []
            if c_date:
                date_parts.append(f"issued {c_date}")
            if c_expiry:
                date_parts.append(f"expires {c_expiry}")
            if date_parts:
                r3 = p.add_run(f"  ({', '.join(date_parts)})")
                _set_run_font(r3, base - 0.5, palette["muted"])
            if c_cred_id:
                r4 = p.add_run(f"  ID: {c_cred_id}")
                _set_run_font(r4, base - 0.5, palette["muted"])
            _item_gap(doc)

    # ── Publications ─────────────────────────────────────────────────────
    pubs = resume.get("publications") or []
    if pubs:
        _section(doc, "Publications", palette, base)
        for pub in pubs:
            title = _safe(pub.get("title"))
            venue = _safe(pub.get("venue"))
            d = _safe(pub.get("date"))
            pub_type = _safe(pub.get("type", ""))
            url_raw = _safe(pub.get("url", ""))
            url = _normalize_url(url_raw) if url_raw else None
            authors = ", ".join(_safe(a) for a in (pub.get("authors") or []) if _safe(a))
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=2, line=1.4)
            if url:
                _add_hyperlink(p, url, title, palette["accent"], base - 0.5)
            else:
                r = p.add_run(title)
                _set_run_font(r, base - 0.5, palette["body"], bold=True)
            if pub_type and f"[{pub_type.lower()}]" not in title.lower():
                rt = p.add_run(f"  [{pub_type}]")
                _set_run_font(rt, base - 0.5, palette["muted"])
            if venue:
                r2 = p.add_run(f"   -   {venue}")
                _set_run_font(r2, base - 0.5, palette["body"], italic=True)
            if d:
                r3 = p.add_run(f"  ({d})")
                _set_run_font(r3, base - 0.5, palette["muted"])
            if authors:
                pa = doc.add_paragraph()
                _set_paragraph_spacing(pa, after=0, line=1.4)
                ra = pa.add_run(authors)
                _set_run_font(ra, base - 1.5, palette["muted"])
            _item_gap(doc)

    # ── Patents ──────────────────────────────────────────────────────────
    patents = resume.get("patents") or []
    if patents:
        _section(doc, "Patents", palette, base)
        for pt in patents:
            pt_title = _safe(pt.get("title"))
            pt_url_raw = _safe(pt.get("url", ""))
            pt_url = _normalize_url(pt_url_raw) if pt_url_raw else None
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=2, line=1.4)
            if pt_url:
                _add_hyperlink(p, pt_url, pt_title, palette["accent"], base - 0.5)
            else:
                r = p.add_run(pt_title)
                _set_run_font(r, base - 0.5, palette["body"], bold=True)
            num = _safe(pt.get("patent_number"))
            d = _safe(pt.get("date"))
            if num:
                r2 = p.add_run(f"   -   {num}")
                _set_run_font(r2, base - 0.5, palette["body"])
            if d:
                r3 = p.add_run(f"  ({d})")
                _set_run_font(r3, base - 0.5, palette["muted"])
            desc = _safe(pt.get("description"))
            if desc:
                p2 = doc.add_paragraph()
                _set_paragraph_spacing(p2, after=0, line=1.4)
                r4 = p2.add_run(desc)
                _set_run_font(r4, base - 1.5, palette["muted"])
            _item_gap(doc)

    # ── Awards ───────────────────────────────────────────────────────────
    awards = resume.get("awards_and_honors") or []
    if awards:
        _section(doc, "Awards & Honors", palette, base)
        for a in awards:
            v = _safe(a) if not isinstance(a, dict) else _safe(a.get("name") or a.get("title"))
            if v:
                _bullet(doc, v, palette, base)

    # ── Volunteer ────────────────────────────────────────────────────────
    vols = resume.get("volunteer_experience") or []
    if vols:
        _section(doc, "Volunteer Experience", palette, base)
        for v in vols:
            org = _safe(v.get("organization"))
            role = _safe(v.get("role"))
            sd, ed = _safe(v.get("start_date")), _safe(v.get("end_date"))
            dates_str = f"{sd}  -  {ed}".strip(" -") if (sd or ed) else ""

            # Mirror Experience / Education: title + dates share one row.
            def _fill_left(lp, _role=role, _org=org):
                r = lp.add_run(_role)
                _set_run_font(r, base, palette["body"], bold=True)
                if _org:
                    r2 = lp.add_run(f"   -   {_org}")
                    _set_run_font(r2, base - 0.5, palette["body"])

            def _fill_right(rp, _dates=dates_str):
                if _dates:
                    rr = rp.add_run(_dates)
                    _set_run_font(rr, base - 1.5, palette["muted"])

            _two_col_row_custom(doc, _fill_left, _fill_right, after_pt=0)
            desc = _safe(v.get("description"))
            if desc:
                p3 = doc.add_paragraph()
                _set_paragraph_spacing(p3, after=0, line=1.4)
                r4 = p3.add_run(desc)
                _set_run_font(r4, base - 1.5, palette["muted"])
            _item_gap(doc)

    # ── Languages ────────────────────────────────────────────────────────
    langs = resume.get("languages") or []
    if langs and len(langs) > 1:
        _section(doc, "Languages", palette, base)
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, after=2, line=1.4)
        for i, l in enumerate(langs):
            if i > 0:
                sep = p.add_run("  ·  ")
                _set_run_font(sep, base - 0.5, palette["muted"])
            lang = _safe(l.get("language"))
            prof = _safe(l.get("proficiency"))
            r = p.add_run(lang)
            _set_run_font(r, base - 0.5, palette["body"], bold=True)
            if prof:
                r2 = p.add_run(f" ({prof})")
                _set_run_font(r2, base - 0.5, palette["muted"])

    # ── Interests ────────────────────────────────────────────────────────
    interests = [_safe(i) for i in (resume.get("interests") or []) if _safe(i)]
    if interests:
        _section(doc, "Interests", palette, base)
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, after=2, line=1.4)
        r = p.add_run(", ".join(interests))
        _set_run_font(r, base - 0.5, palette["body"])

    # ── References ───────────────────────────────────────────────────────
    references = resume.get("references")
    if references:
        _section(doc, "References", palette, base)
        if isinstance(references, str):
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=2, line=1.4)
            r = p.add_run(_safe(references))
            _set_run_font(r, base - 0.5, palette["body"])
        elif isinstance(references, list):
            for ref in references:
                ref_str = _safe(ref)
                if ref_str:
                    _bullet(doc, ref_str, palette, base)

    return doc


def _apply_docx_background(doc: Document, hex_no_hash: str) -> None:
    """Set the document page background colour in a DOCX file."""
    color = hex_no_hash.upper()
    doc_element = doc.element
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), color)
    body = doc_element.body
    doc_element.insert(list(doc_element).index(body), bg)
    settings = doc.settings.element
    disp = OxmlElement("w:displayBackgroundShape")
    settings.insert(0, disp)



def generate_docx_resume(
    resume: dict,
    template: TemplateName = "classic_ats",
    font_size: FontSize = "normal",
) -> bytes:
    """Generate an ATS-friendly DOCX resume from resume JSON. Returns DOCX bytes."""
    from app.documents._normalize import normalize_resume, resolve_base_font_size
    resume = normalize_resume(resume)
    base = resolve_base_font_size(font_size)

    # latex_serif gets the dedicated Computer-Modern layout with embedded fonts,
    # so the DOCX matches the PDF on any machine.
    if template == "latex_serif":
        doc = _build_resume_doc_serif(resume, base)
        buf = io.BytesIO()
        doc.save(buf)
        try:
            return _embed_fonts_docx(buf.getvalue())
        except Exception:
            return buf.getvalue()  # fall back to font-name-only if embedding fails

    doc = _build_resume_doc(resume, template, base)
    palette = PALETTES_HEX.get(template, PALETTES_HEX["classic_ats"])
    if "bg" in palette:
        _apply_docx_background(doc, palette["bg"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_docx_cover_letter(
    cover_letter: str,
    resume: dict,
    template: TemplateName = "modern_clean",
    font_size: FontSize = "normal",
    hiring_manager: str | None = None,
    company_name: str | None = None,
    date_str: str | None = None,
) -> bytes:
    """Generate a one-page cover letter DOCX matching the PDF layout."""
    from app.documents._normalize import resolve_base_font_size
    palette = PALETTES_HEX.get(template, PALETTES_HEX["modern_clean"])
    base = resolve_base_font_size(font_size)
    pi = (resume or {}).get("personal_info", {}) or {}
    name = _safe(pi.get("full_name")) or "Your Name"

    doc = Document()
    # Match the PDF cover-letter generator: 0.85in all sides.
    _set_doc_margins(doc, side_cm=2.159, vert_cm=2.159)

    # ── Sender block (right-aligned) ─────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(p, after=2, line=1.15)
    r = p.add_run(name)
    _set_run_font(r, base + 4, palette["name"], bold=True)

    for key in ("email", "phone", "location", "linkedin", "github", "website", "portfolio"):
        v = _safe(pi.get(key))
        if not v:
            continue
        display = v.replace("https://", "").replace("http://", "").rstrip("/") if key in ("linkedin", "github", "website", "portfolio") else v
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_spacing(sp, after=1, line=1.5)
        sr = sp.add_run(display)
        _set_run_font(sr, base - 1.5, palette["muted"])

    if template != "classic_ats":
        rule_p = doc.add_paragraph()
        rule_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_spacing(rule_p, before=4, after=4)
        _bottom_border(rule_p, palette["accent"], size=8)

    # ── Date ─────────────────────────────────────────────────────────────
    display_date = date_str if date_str else date.today().strftime("%B %d, %Y")
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=18, after=14)
    r = p.add_run(display_date)
    _set_run_font(r, base - 0.5, palette["muted"])

    # ── Recipient ────────────────────────────────────────────────────────
    recip_lines = []
    if hiring_manager and hiring_manager.lower() not in ("hiring team", "hiring manager", ""):
        recip_lines.append(hiring_manager)
    if company_name:
        recip_lines.append(company_name)
    for line in recip_lines:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, after=2, line=1.45)
        r = p.add_run(line)
        _set_run_font(r, base - 0.5, palette["body"])
    if recip_lines:
        # trailing space
        doc.add_paragraph()

    # ── Salutation ───────────────────────────────────────────────────────
    salutation = "Dear Hiring Team,"
    if hiring_manager and hiring_manager.lower() not in ("hiring team", ""):
        first = hiring_manager.strip().split()[0]
        if first and first[0].isalpha():
            salutation = f"Dear {first},"
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, after=10)
    r = p.add_run(salutation)
    _set_run_font(r, base, palette["body"])

    # ── Body ─────────────────────────────────────────────────────────────
    text = (cover_letter or "").strip()
    paras = [pp.strip() for pp in text.split("\n\n") if pp.strip()]
    for para in paras:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_paragraph_spacing(p, after=10, line=1.55)
        for i, chunk in enumerate(para.split("\n")):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(chunk)
            _set_run_font(r, base, palette["body"])

    # ── Signature ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=10, after=4)
    r = p.add_run("Sincerely,")
    _set_run_font(r, base, palette["body"])

    doc.add_paragraph()  # signature gap

    p = doc.add_paragraph()
    r = p.add_run(name)
    _set_run_font(r, base, palette["name"], bold=True)

    if "bg" in palette:
        _apply_docx_background(doc, palette["bg"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

"""python-docx output — every template × font size combo."""
from __future__ import annotations

import io

import pytest
from docx import Document

pytestmark = pytest.mark.documents


def _extract_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


@pytest.mark.parametrize("template", ["classic_ats", "modern_clean", "executive_dark", "dark_theme"])
@pytest.mark.parametrize("font_size", ["small", "normal", "large"])
def test_resume_docx_renders_all_combos(sample_resume, template, font_size):
    from app.documents.docx_generator import generate_docx_resume

    data = generate_docx_resume(sample_resume, template=template, font_size=font_size)
    # DOCX is a ZIP archive — first bytes are "PK".
    assert data[:2] == b"PK"
    assert len(data) > 5_000
    text = _extract_text(data)
    assert "Ada Lovelace" in text


def test_docx_includes_summary(sample_resume):
    from app.documents.docx_generator import generate_docx_resume

    data = generate_docx_resume(sample_resume)
    text = _extract_text(data)
    assert "backend engineer" in text.lower()


def test_docx_minimal(minimal_resume):
    from app.documents.docx_generator import generate_docx_resume

    data = generate_docx_resume(minimal_resume)
    assert data[:2] == b"PK"
    text = _extract_text(data)
    assert "Grace Hopper" in text


@pytest.mark.parametrize("font_size", ["small", "normal", "large"])
def test_docx_latex_serif_embeds_fonts(sample_resume, font_size):
    """latex_serif DOCX mirrors the PDF: Computer Modern + Font Awesome embedded
    into the package so it renders the same on any machine."""
    import zipfile

    from app.documents.docx_generator import generate_docx_resume

    data = generate_docx_resume(sample_resume, template="latex_serif", font_size=font_size)
    assert data[:2] == b"PK"
    assert "Ada Lovelace" in _extract_text(data)
    z = zipfile.ZipFile(io.BytesIO(data))
    assert any("word/fonts/font" in n for n in z.namelist()), "fonts not embedded"
    ft = z.read("word/fontTable.xml").decode()
    assert "CMU Serif" in ft and "SR FA Solid" in ft and "SR FA Brands" in ft


def test_docx_latex_serif_minimal(minimal_resume):
    from app.documents.docx_generator import generate_docx_resume

    data = generate_docx_resume(minimal_resume, template="latex_serif")
    assert data[:2] == b"PK"
    assert "Grace Hopper" in _extract_text(data)


def test_docx_cover_letter(sample_resume):
    from app.documents.docx_generator import generate_docx_cover_letter

    body = "Hi team,\n\nI am applying for the Senior Backend role."
    data = generate_docx_cover_letter(
        cover_letter=body,
        resume=sample_resume,
        template="modern_clean",
    )
    assert data[:2] == b"PK"
    text = _extract_text(data)
    assert "Ada Lovelace" in text
    assert "applying" in text.lower()


def test_docx_cover_letter_dark_theme_has_background(sample_resume):
    """Dark-theme cover letters must embed the page background colour or
    the light run colours render invisibly on Word's default white page."""
    import zipfile, io as _io
    from app.documents.docx_generator import generate_docx_cover_letter

    data = generate_docx_cover_letter(
        cover_letter="Hi team.",
        resume=sample_resume,
        template="dark_theme",
    )
    assert data[:2] == b"PK"
    with zipfile.ZipFile(_io.BytesIO(data)) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")
    # ``<w:background w:color="0D0F1A">`` is injected by _apply_docx_background.
    assert "w:background" in document_xml
    assert "0D0F1A" in document_xml.upper()

"""Tiny binary samples for upload-endpoint tests.

These are intentionally minimal — just enough for pypdf / python-docx to
parse without complaint. Constructed in-memory so the test suite has no
external file dependencies.
"""
from __future__ import annotations

import io
import json


def make_text_bytes(text: str = "Hello resume world.") -> bytes:
    return text.encode("utf-8")


def make_json_resume_bytes(resume: dict | None = None) -> bytes:
    resume = resume or {
        "personal_info": {"full_name": "Linus Torvalds", "email": "linus@example.com"},
        "professional_summary": "Kernel maintainer.",
    }
    return json.dumps(resume).encode("utf-8")


def make_pdf_bytes(text: str = "Hello PDF resume.") -> bytes:
    """Build a real PDF that pypdf can extract text from."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica", 12)
    c.drawString(100, 750, text)
    c.showPage()
    c.save()
    return buf.getvalue()


def make_docx_bytes(text: str = "Hello DOCX resume.") -> bytes:
    """Build a real DOCX so python-docx can round-trip the text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    doc.add_paragraph("Second line.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
